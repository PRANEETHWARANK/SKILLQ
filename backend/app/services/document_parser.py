import io
import re
from typing import Dict, Any
from pypdf import PdfReader
import docx

class DocumentParserService:
    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
        ext = filename.lower().split('.')[-1]
        text = ""
        
        if ext == "pdf":
            try:
                reader = PdfReader(io.BytesIO(file_bytes))
                pages_text = []
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        pages_text.append(extracted)
                text = "\n".join(pages_text)
            except Exception as e:
                raise ValueError(f"Failed to parse PDF: {str(e)}")
                
        elif ext in ["docx", "doc"]:
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            paragraphs.append(row_text)
                text = "\n".join(paragraphs)
            except Exception as e:
                raise ValueError(f"Failed to parse DOCX: {str(e)}")
                
        elif ext in ["txt", "md"]:
            try:
                text = file_bytes.decode("utf-8", errors="replace")
            except Exception as e:
                raise ValueError(f"Failed to parse text file: {str(e)}")
        else:
            raise ValueError(f"Unsupported file format: .{ext}. Please upload a PDF, DOCX, or TXT file.")
            
        cleaned = DocumentParserService.clean_text(text)
        if not cleaned or len(cleaned.strip()) < 30:
            raise ValueError("Extracted text is empty or too short to be a valid resume.")
        return cleaned

    @staticmethod
    def clean_text(text: str) -> str:
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        return text.strip()

    @staticmethod
    def segment_sections(text: str) -> Dict[str, str]:
        sections = {
            "contact": "",
            "summary": "",
            "skills": "",
            "experience": "",
            "projects": "",
            "education": "",
            "certifications": "",
            "other": ""
        }
        
        headers = [
            ("skills", r"(?i)^(technical\s+skills|skills\s*&\s*abilities|skills|tech\s+stack|core\s+competencies)"),
            ("experience", r"(?i)^(work\s+experience|professional\s+experience|experience|employment\s+history|internships?)"),
            ("projects", r"(?i)^(projects|academic\s+projects|personal\s+projects|key\s+projects)"),
            ("education", r"(?i)^(education|academic\s+background|qualifications)"),
            ("certifications", r"(?i)^(certifications|licenses|courses|achievements|awards)"),
            ("summary", r"(?i)^(professional\s+summary|summary|profile|about\s+me|objective)")
        ]
        
        lines = text.split("\n")
        current_section = "summary"
        
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
                
            matched_header = False
            for sec_name, pattern in headers:
                if re.match(pattern, line_str) and len(line_str) < 40:
                    current_section = sec_name
                    matched_header = True
                    break
                    
            if not matched_header:
                sections[current_section] = sections[current_section] + "\n" + line_str
                
        return {k: v.strip() for k, v in sections.items() if v.strip()}
